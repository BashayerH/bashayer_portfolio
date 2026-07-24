#!/usr/bin/env python3
"""Generate a Word (.docx) version of Bashayer Alhawsawi's portfolio.

Content mirrors portfolio-cv.html. Images are pulled from img/ (webp is
converted to png and oversized images are downscaled to keep the file small).

Usage: python3 scripts/generate_docx.py
Output: Bashayer-Alhawsawi-Portfolio.docx (repo root)
"""

import os
import tempfile

from PIL import Image

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG_DIR = os.path.join(ROOT, "img")
OUT = os.path.join(ROOT, "Bashayer-Alhawsawi-Portfolio.docx")

GOLD = RGBColor(0xC9, 0x93, 0x3A)
DARK = RGBColor(0x2A, 0x25, 0x1E)
GREY = RGBColor(0x5A, 0x53, 0x4A)
LINK = "0563C1"

_tmp = tempfile.mkdtemp(prefix="docx_img_")


def prep_image(name, max_px=1000):
    """Return a path usable by python-docx (png/jpg), downscaled if large."""
    src = os.path.join(IMG_DIR, name)
    im = Image.open(src)
    if im.mode in ("RGBA", "P", "LA"):
        im = im.convert("RGB")
    im.thumbnail((max_px, max_px), Image.LANCZOS)
    out = os.path.join(_tmp, os.path.splitext(name)[0].replace("/", "_") + ".jpg")
    im.save(out, "JPEG", quality=82)
    return out


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK)
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def section_heading(doc, num, title):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    r = p.add_run(f"{num}  ")
    r.font.color.rgb = GOLD
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.bold = True
    r = p.add_run(title.upper())
    r.font.color.rgb = GOLD
    r.font.size = Pt(10)
    r.bold = True
    r.font.name = "Consolas"
    h = doc.add_paragraph()
    hr = h.add_run("")
    # thin gold rule under heading
    pborder = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9933A")
    pborder.append(bottom)
    h.paragraph_format.element.get_or_add_pPr().append(pborder)
    h.paragraph_format.space_after = Pt(8)


def images_row(doc, images, total_width=6.6):
    n = len(images)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_w = total_width / n
    for i, name in enumerate(images):
        cell = table.rows[0].cells[i]
        cell.width = Inches(cell_w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(prep_image(name), width=Inches(cell_w - 0.12))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def project(doc, index, kind, title, category, desc, result, tags, images, links):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = DARK
    r = p.add_run(f"    {index} / {kind}")
    r.font.color.rgb = GOLD
    r.font.size = Pt(9)
    r.font.name = "Consolas"

    c = doc.add_paragraph()
    cr = c.add_run(category.upper())
    cr.font.color.rgb = GOLD
    cr.font.size = Pt(8.5)
    cr.bold = True
    c.paragraph_format.space_after = Pt(4)

    d = doc.add_paragraph()
    dr = d.add_run(desc)
    dr.font.size = Pt(10.5)
    dr.font.color.rgb = DARK

    rp = doc.add_paragraph()
    rr = rp.add_run("\u2605 " + result)
    rr.font.size = Pt(10)
    rr.italic = True
    rr.font.color.rgb = RGBColor(0x9A, 0x6E, 0x22)
    rp.paragraph_format.space_after = Pt(6)

    images_row(doc, images)

    tp = doc.add_paragraph()
    tr = tp.add_run("Tech:  ")
    tr.bold = True
    tr.font.size = Pt(9.5)
    tr.font.color.rgb = GREY
    tr2 = tp.add_run("  \u00b7  ".join(tags))
    tr2.font.size = Pt(9.5)
    tr2.font.color.rgb = DARK

    lp = doc.add_paragraph()
    lr = lp.add_run("Links:  ")
    lr.bold = True
    lr.font.size = Pt(9.5)
    lr.font.color.rgb = GREY
    for i, (label, url) in enumerate(links):
        if i:
            lp.add_run("    ")
        add_hyperlink(lp, url, label)
    lp.paragraph_format.space_after = Pt(10)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # ---- Cover ----
    kicker = doc.add_paragraph()
    kr = kicker.add_run("PORTFOLIO \u00b7 2025")
    kr.font.color.rgb = GOLD
    kr.font.name = "Consolas"
    kr.font.size = Pt(10)
    kr.bold = True

    name = doc.add_paragraph()
    nr = name.add_run("Bashayer Alhawsawi")
    nr.bold = True
    nr.font.size = Pt(34)
    nr.font.color.rgb = DARK
    name.paragraph_format.space_after = Pt(2)

    role = doc.add_paragraph()
    rr = role.add_run("Mobile Developer  \u00b7  Android & Cross-Platform")
    rr.font.size = Pt(14)
    rr.font.color.rgb = GOLD
    rr.bold = True

    tag = doc.add_paragraph()
    tr = tag.add_run(
        "I build mobile apps that people actually use \u2014 native Android with Kotlin and "
        "cross-platform with Flutter, engineered to feel effortless. Based in Saudi Arabia."
    )
    tr.font.size = Pt(11)
    tr.font.color.rgb = GREY
    tag.paragraph_format.space_after = Pt(8)

    stats = doc.add_paragraph()
    for label in ["4+ Years Experience", "5 Shipped Projects", "3 Apps on Google Play"]:
        sr = stats.add_run("   \u2b29 " + label + "   ")
        sr.font.size = Pt(10.5)
        sr.bold = True
        sr.font.color.rgb = DARK

    contact = doc.add_paragraph()
    contact.paragraph_format.space_before = Pt(6)
    parts = [
        ("Email", "bashayer.eh.19@gmail.com", "mailto:bashayer.eh.19@gmail.com"),
        ("WhatsApp", "+966 54 483 0939", "https://wa.me/966544830939"),
        ("LinkedIn", "linkedin.com/in/bashayereh", "https://www.linkedin.com/in/bashayereh"),
        ("GitHub", "github.com/BashayerH", "https://github.com/BashayerH"),
        ("X", "x.com/fo0oshaty", "https://x.com/fo0oshaty"),
    ]
    for i, (k, label, url) in enumerate(parts):
        if i:
            contact.add_run("    |    ").font.color.rgb = GOLD
        kr = contact.add_run(k + ": ")
        kr.bold = True
        kr.font.size = Pt(9.5)
        kr.font.color.rgb = GOLD
        add_hyperlink(contact, url, label)

    # ---- About ----
    section_heading(doc, "01", "About \u2014 How I Became a Mobile Developer")
    about_paras = [
        "My journey began at the Tuwaiq Academy bootcamp, where I built my first complete "
        "application and truly understood the fundamentals of programming.",
        "After graduating, I worked on small projects that strengthened my skills and gave me "
        "real experience solving practical problems through code.",
        "My first professional role introduced me to Flutter, where I spent a year and a half "
        "building cross-platform mobile applications.",
        "Later, I returned to my foundation: Kotlin. Since then I've been continuously refining "
        "my Android expertise, improving my architecture skills, and building applications that "
        "are reliable, scalable, and designed to last.",
    ]
    for t in about_paras:
        p = doc.add_paragraph()
        r = p.add_run(t)
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK

    # ---- Skills ----
    section_heading(doc, "02", "Tools I Trust")
    skills = [
        ("Kotlin", "Coroutines, Flow, Jetpack Compose, MVVM, Room, Retrofit \u2014 native Android from architecture to animation."),
        ("Flutter", "Cross-platform UI at speed. Bloc/Provider state management, custom animations, native feel on both platforms."),
        ("Clean Architecture", "Domain, Data and Presentation layers. Testable, maintainable, built to survive the next developer."),
        ("Maps SDK", "Live tracking, geofencing, route calculation, custom overlays \u2014 real-time location at production scale."),
    ]
    for name_, desc in skills:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(name_ + " \u2014 ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = GREY

    # ---- Projects ----
    section_heading(doc, "03", "Selected Work \u2014 My Projects")

    project(
        doc, "01", "Web", "Ma'an Initiative", "Web Platform \u00b7 Healthcare Education",
        "Under the Ma'an Initiative, I built a web platform serving the healthcare education "
        "sector \u2014 a responsive, accessible site delivering educational content to its audience.",
        "Live platform \u00b7 Healthcare education sector",
        ["HTML5", "CSS3", "JavaScript", "Responsive"],
        ["maan_5.png", "maan_1.png", "maan03.png"],
        [("Visit Website", "https://maan-bfd2f.web.app")],
    )
    project(
        doc, "02", "Mobile", "Minute \u2014 Driver App", "Real-Time Ride-Hailing",
        "Drivers needed real-time reliability. I built a production-grade app with Google Maps, "
        "live ride handling, and earnings tracking \u2014 performance-first, always.",
        "Live on Google Play \u00b7 Real-time ride-hailing",
        ["Flutter", "Kotlin", "Firebase", "Google Maps"],
        ["driver01.jpg", "driver02.png", "driver03.jpg"],
        [("Google Play", "https://play.google.com/store/apps/details?id=com.taxi.minutedriver"),
         ("Demo", "https://drive.google.com/file/d/1cHaCTntuuKOEfCkGW130yeb3p_Zv1gYy/view?usp=share_link")],
    )
    project(
        doc, "03", "Mobile", "Minute \u2014 Client App", "Production-Ready Passenger App",
        "A ride-hailing app that lets users request trips, track drivers live, and reach "
        "destinations safely and quickly \u2014 with intuitive booking flows, push notifications, "
        "and a UX people trust.",
        "Live on Google Play \u00b7 Production-ready passenger app",
        ["Kotlin", "Flutter", "Firebase", "Google Maps"],
        ["clinet01.png", "clinet02.png", "clinet03.png"],
        [("Google Play", "https://play.google.com/store/apps/details?id=com.taxi.minute"),
         ("Demo", "https://drive.google.com/file/d/1zgE6lEMfZeEql-GFmq3ByigTSRVs9H5-/view?usp=share_link")],
    )
    project(
        doc, "04", "Mobile", "Minhaj App", "Religious & Educational Platform",
        "An integrated digital platform for religious and educational services, developed under "
        "the Elm company as part of the National Skills Development Program.",
        "Live on Google Play \u00b7 National Skills Development Program",
        ["Kotlin", "Android", "Firebase", "MVVM", "Jetpack Compose"],
        ["minhaj01.jpg", "minhaj02.jpg", "minhaj03.jpg"],
        [("Google Play", "https://play.google.com/store/apps/details?id=sa.elm.mwl&hl=ar")],
    )
    project(
        doc, "05", "KMP", "Meme Editor",
        "Kotlin Multiplatform \u00b7 Android \u00b7 iOS \u00b7 Web \u00b7 Desktop",
        "A Kotlin Multiplatform app targeting Android, iOS, Web, and Desktop, letting users "
        "create and customize memes with text editing and sharing. Built to explore modern "
        "cross-platform development from one codebase.",
        "Exploration project \u00b7 Kotlin Multiplatform",
        ["Kotlin Multiplatform", "Compose", "Android", "iOS", "Web", "Desktop"],
        ["memeE.webp", "memeEdit.webp"],
        [("GitHub", "https://github.com/BashayerH")],
    )

    # ---- Contact ----
    section_heading(doc, "04", "Let's Talk")
    p = doc.add_paragraph()
    r = p.add_run("If you're building something ambitious, let's talk.")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = DARK
    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        "Available for new challenges. I take on selective work, which means what we build gets "
        "my full attention. Typically replies within 24 hours \u2014 no automated responses."
    )
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GREY

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(10)
    fr = foot.add_run("\u00a9 2025 Bashayer Alhawsawi \u00b7 Built with intention.")
    fr.font.size = Pt(9)
    fr.italic = True
    fr.font.color.rgb = GREY

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    main()
